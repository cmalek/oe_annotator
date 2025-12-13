"""Unit tests for DOCXExporter."""

import tempfile
from pathlib import Path

import pytest
from docx import Document

from oeapp.services.export_docx import DOCXExporter
from tests.conftest import create_test_project, create_test_sentence


class TestDOCXExporter:
    """Test cases for DOCXExporter."""

    def test_export_creates_document(self, db_session, tmp_path):
        """Test export() creates a DOCX file."""
        project = create_test_project(db_session, name="Test Project", text="Se cyning.")
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        result = exporter.export(project.id, output_path)

        assert result is True
        assert output_path.exists()

    def test_export_returns_false_when_project_not_found(self, db_session, tmp_path):
        """Test export() returns False when project doesn't exist."""
        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        result = exporter.export(99999, output_path)

        assert result is False

    def test_export_includes_project_title(self, db_session, tmp_path):
        """Test export() includes project name as heading."""
        project = create_test_project(db_session, name="My Test Project", text="")
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # First paragraph should be the heading
        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].text == "My Test Project"

    def test_export_includes_sentence_numbers(self, db_session, tmp_path):
        """Test export() includes paragraph and sentence numbers."""
        project = create_test_project(db_session, name="Test", text="Se cyning. Þæt scip.")
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        # Should contain paragraph and sentence number markers
        assert "¶[" in text
        assert "S[" in text

    def test_export_includes_paragraph_breaks(self, db_session, tmp_path):
        """Test export() adds extra blank lines for paragraph starts."""
        project = create_test_project(db_session, name="Test", text="")
        db_session.commit()

        # Create sentences with paragraph breaks
        sentence1 = create_test_sentence(
            db_session, project_id=project.id, text="First paragraph.", display_order=1, is_paragraph_start=True
        )
        sentence2 = create_test_sentence(
            db_session, project_id=project.id, text="Second paragraph.", display_order=2, is_paragraph_start=True
        )
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Should have multiple paragraphs (title + sentences + breaks)
        assert len(doc.paragraphs) > 2

    def test_export_includes_translation(self, db_session, tmp_path):
        """Test export() includes modern translation when available."""
        project = create_test_project(db_session, name="Test", text="Se cyning.")
        db_session.commit()

        # Get the sentence and add translation
        sentence = project.sentences[0]
        sentence.text_modern = "The king"
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "The king" in text

    def test_export_handles_missing_translation(self, db_session, tmp_path):
        """Test export() handles sentences without translation."""
        project = create_test_project(db_session, name="Test", text="Se cyning.")
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Should still create document successfully
        assert len(doc.paragraphs) > 0

    def test_export_with_annotations_includes_superscripts(self, db_session, tmp_path):
        """Test export() includes superscript POS annotations."""
        project = create_test_project(db_session, name="Test", text="Se cyning")
        db_session.commit()

        sentence = project.sentences[0]
        token = sentence.tokens[0]

        # Create annotation with POS
        if token.annotation:
            annotation = token.annotation
        else:
            from oeapp.models.annotation import Annotation
            annotation = Annotation(token_id=token.id)
            db_session.add(annotation)
            db_session.flush()

        annotation.pos = "R"
        annotation.pronoun_type = "d"
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Find paragraph with Old English text
        oe_para = None
        for para in doc.paragraphs:
            if any("Se" in run.text for run in para.runs):
                oe_para = para
                break

        assert oe_para is not None
        # Check for superscript runs
        has_superscript = any(run.font.superscript for run in oe_para.runs)
        assert has_superscript

    def test_export_with_annotations_includes_subscripts(self, db_session, tmp_path):
        """Test export() includes subscript gender/context annotations."""
        project = create_test_project(db_session, name="Test", text="Se cyning")
        db_session.commit()

        sentence = project.sentences[0]
        token = sentence.tokens[0]

        # Create annotation with gender
        if token.annotation:
            annotation = token.annotation
        else:
            from oeapp.models.annotation import Annotation
            annotation = Annotation(token_id=token.id)
            db_session.add(annotation)
            db_session.flush()

        annotation.pos = "N"
        annotation.gender = "m"
        annotation.case = "n"
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        # Find paragraph with Old English text
        oe_para = None
        for para in doc.paragraphs:
            if any("cyning" in run.text for run in para.runs):
                oe_para = para
                break

        assert oe_para is not None
        # Check for subscript runs
        has_subscript = any(run.font.subscript for run in oe_para.runs)
        assert has_subscript

    def test_export_with_notes_includes_notes(self, db_session, tmp_path):
        """Test export() includes notes for sentences."""
        project = create_test_project(db_session, name="Test", text="Se cyning")
        db_session.commit()

        sentence = project.sentences[0]
        token = sentence.tokens[0]

        # Create a note
        from oeapp.models.note import Note
        note = Note(
            sentence_id=sentence.id,
            start_token=token.id,
            end_token=token.id,
            note_text_md="This is a test note"
        )
        db_session.add(note)
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        assert "This is a test note" in text
        assert "1." in text  # Note numbering

    def test_export_with_multiple_notes_orders_correctly(self, db_session, tmp_path):
        """Test export() orders multiple notes by token position."""
        project = create_test_project(db_session, name="Test", text="Se cyning fēoll")
        db_session.commit()

        sentence = project.sentences[0]
        tokens = list(sentence.tokens)

        # Create notes in reverse order
        from oeapp.models.note import Note
        note2 = Note(
            sentence_id=sentence.id,
            start_token=tokens[2].id,
            end_token=tokens[2].id,
            note_text_md="Note on third token"
        )
        note1 = Note(
            sentence_id=sentence.id,
            start_token=tokens[0].id,
            end_token=tokens[0].id,
            note_text_md="Note on first token"
        )
        db_session.add(note2)
        db_session.add(note1)
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        exporter.export(project.id, output_path)

        doc = Document(str(output_path))
        text = "\n".join([para.text for para in doc.paragraphs])

        # First note should appear before second note
        first_pos = text.find("Note on first token")
        second_pos = text.find("Note on third token")
        assert first_pos < second_pos

    def test_export_empty_project_creates_document(self, db_session, tmp_path):
        """Test export() creates document even for empty project."""
        project = create_test_project(db_session, name="Empty Project", text="")
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        result = exporter.export(project.id, output_path)

        assert result is True
        assert output_path.exists()

        doc = Document(str(output_path))
        # Should have at least the title
        assert len(doc.paragraphs) > 0
        assert doc.paragraphs[0].text == "Empty Project"

    def test_setup_document_styles_sets_margins(self, db_session):
        """Test _setup_document_styles() sets document margins."""
        from docx import Document as DocxDocument

        exporter = DOCXExporter(db_session)
        doc = DocxDocument()

        exporter._setup_document_styles(doc)

        section = doc.sections[0]
        assert section.top_margin.inches == 1.0
        assert section.left_margin.inches == 1.0
        assert section.right_margin.inches == 1.0
        assert section.bottom_margin.inches == 1.0

    def test_export_handles_sentence_without_tokens(self, db_session, tmp_path):
        """Test export() handles sentence with no tokens gracefully."""
        project = create_test_project(db_session, name="Test", text="")
        db_session.commit()

        # Create sentence manually without tokens
        from oeapp.models.sentence import Sentence
        sentence = Sentence(
            project_id=project.id,
            display_order=1,
            text_oe="Test sentence",
            paragraph_number=1,
            sentence_number_in_paragraph=1
        )
        db_session.add(sentence)
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        result = exporter.export(project.id, output_path)

        assert result is True
        assert output_path.exists()

    def test_export_handles_file_write_error(self, db_session, tmp_path, monkeypatch):
        """Test export() handles file write errors gracefully."""
        project = create_test_project(db_session, name="Test", text="Se cyning.")
        db_session.commit()

        exporter = DOCXExporter(db_session)
        output_path = tmp_path / "test.docx"

        # Mock Document.save to raise OSError
        original_save = None
        from docx.document import Document as DocumentClass

        def mock_save(self, path):
            raise OSError("Permission denied")

        # Patch the save method on the instance
        monkeypatch.setattr(DocumentClass, "save", mock_save)

        result = exporter.export(project.id, output_path)

        assert result is False

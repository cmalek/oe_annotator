"""DOCX export service for Ænglisc Toolkit."""

from typing import TYPE_CHECKING

from docx import Document
from docx.shared import Pt

from oeapp.mixins import AnnotationTextualMixin
from oeapp.models.project import Project

if TYPE_CHECKING:
    from pathlib import Path

    from docx.document import Document as DocumentObject
    from sqlalchemy.orm import Session

    from oeapp.models.annotation import Annotation
    from oeapp.models.sentence import Sentence


class DOCXExporter(AnnotationTextualMixin):
    """
    Exports annotated Old English text to DOCX format.

    Args:
        session: SQLAlchemy session

    """

    def __init__(self, session: Session) -> None:
        """
        Initialize exporter.

        Args:
            session: SQLAlchemy session

        """
        self.session = session

    def export(self, project_id: int, output_path: Path) -> bool:
        """
        Export project to DOCX file.

        Args:
            project_id: Project ID to export
            output_path: Path to output DOCX file

        Returns:
            True if successful, False otherwise

        """
        doc: DocumentObject = Document()
        self._setup_document_styles(doc)
        project = Project.get(self.session, project_id)
        if project is None:
            return False

        # Add title
        doc.add_heading(project.name, level=1)
        doc.add_paragraph()  # Blank line after title

        for sentence in project.sentences:
            display_order = sentence.display_order
            text_modern = sentence.text_modern

            # Add sentence number
            sentence_num_para = doc.add_paragraph()
            sentence_num_run = sentence_num_para.add_run(f"[{display_order}] ")
            sentence_num_run.bold = True

            # Build Old English sentence with annotations
            self._add_oe_sentence_with_annotations(doc, sentence)

            # Add translation
            if text_modern:
                doc.add_paragraph(text_modern)
            else:
                # Empty translation paragraph
                doc.add_paragraph()

            # Blank line
            doc.add_paragraph()

            # Add notes
            self._add_notes(doc, sentence)

            # Blank line between sentences
            doc.add_paragraph()

        try:
            doc.save(str(output_path))
        except OSError as e:
            print(f"Export error: {e}")  # noqa: T201
            return False
        else:
            return True

    def _setup_document_styles(self, doc: DocumentObject) -> None:
        """
        Set up document styles.

        Args:
            doc: Document to set up

        """
        # Styles are already available in python-docx
        # Title, Body, Default are standard styles

    def _add_oe_sentence_with_annotations(  # noqa: PLR0912, PLR0915
        self,
        doc: DocumentObject,
        sentence: Sentence,
    ) -> None:
        """
        Add Old English sentence with superscript/subscript annotations.

        Args:
            doc: Document to add to
            sentence: Sentence to add

        """
        # Build paragraph with annotations
        para = doc.add_paragraph()  # type: ignore[attr-defined]

        for i, token in enumerate(sentence.tokens):
            surface = token.surface
            annotation = token.annotation

            # Add word
            word_run = para.add_run(surface)
            word_run.italic = True

            if annotation:
                # Build superscript text (POS abbreviations, case, number, gender)
                superscript_parts = []

                # Add POS abbreviations first
                pos_label = self._format_pos(annotation)
                if pos_label:
                    superscript_parts.append(pos_label)
                if annotation.article_type:
                    superscript_parts.append(
                        self.ARTICLE_TYPE_MAP.get(
                            annotation.article_type, annotation.article_type
                        )
                    )

                # Add case/number/gender compact format
                if annotation.case and annotation.number:
                    case_str = self.CASE_MAP.get(annotation.case, "")
                    number_str = self.NUMBER_MAP.get(annotation.number, "")
                    if case_str and number_str:
                        compact = f"{case_str}{number_str}"
                        superscript_parts.append(compact)
                elif annotation.case:
                    superscript_parts.append(
                        self.CASE_MAP.get(annotation.case, annotation.case)
                    )
                elif annotation.number:
                    superscript_parts.append(
                        self.NUMBER_MAP.get(annotation.number, annotation.number)
                    )
                if annotation.gender:
                    superscript_parts.append(annotation.gender)

                # Build subscript text (detailed morphological info)
                subscript_parts = []
                if annotation.declension:
                    subscript_parts.append(annotation.declension)
                if annotation.verb_class:
                    subscript_parts.append(annotation.verb_class)
                # Add compact case/number combinations like "dat1", "acc1"
                if annotation.case and annotation.number:
                    case_str = self.CASE_MAP.get(annotation.case, annotation.case)[:3]
                    number_str = self.NUMBER_MAP.get(
                        annotation.number, annotation.number
                    )
                    if annotation.case in [
                        "d",
                        "a",
                        "g",
                    ]:  # Common cases for compact format
                        compact = f"{case_str}{number_str}"
                        if compact not in subscript_parts:
                            subscript_parts.insert(0, compact)

                # Add uncertain marker
                uncertain_marker = "?" if annotation.uncertain else ""
                if annotation.alternatives_json:
                    alternatives = f" / {annotation.alternatives_json}"
                else:
                    alternatives = ""

                # Add superscript
                if superscript_parts:
                    sup_text = (
                        "".join(superscript_parts) + uncertain_marker + alternatives
                    )
                    sup_run = para.add_run(sup_text)
                    sup_run.font.size = Pt(8)
                    sup_run.font.superscript = True

                # Add subscript
                if subscript_parts:
                    sub_text = "".join(subscript_parts) + uncertain_marker
                    sub_run = para.add_run(sub_text)
                    sub_run.font.size = Pt(8)
                    sub_run.font.subscript = True

            # Add space after word (except last token)
            if i < len(sentence.tokens) - 1:
                para.add_run(" ")

    def _format_pos(self, annotation: Annotation) -> str:
        """
        Format part of speech abbreviation for display.

        Args:
            annotation: Annotation object

        Returns:
            Formatted POS string

        """
        if not annotation.pos:
            return ""

        base = self.PART_OF_SPEECH_MAP.get(annotation.pos, annotation.pos.lower())

        if annotation.pos == "R" and annotation.pronoun_type:
            type_str = self.PRONOUN_TYPE_MAP.get(annotation.pronoun_type, "")
            if type_str:
                return f"{base}{type_str}"
        elif annotation.pos == "V" and annotation.verb_class:
            return f"{base}{annotation.verb_class}"

        return base

    def _add_notes(self, doc: DocumentObject, sentence: Sentence) -> None:
        """
        Add notes for a sentence.

        Notes are sorted by their position in the sentence (by start token
        order_index) and numbered accordingly.

        Args:
            doc: Document to add to
            sentence: Sentence to add notes for

        """
        if not sentence.notes:
            return

        # Sort notes by token position in sentence (earlier tokens = lower numbers)
        notes = self._sort_notes_by_position(sentence)

        # Build token ID to order_index mapping for token lookups
        token_id_to_order: dict[int, int] = {}
        for token in sentence.tokens:
            if token.id:
                token_id_to_order[token.id] = token.order_index

        # Display each note with dynamic numbering (1-based index)
        for note_idx, note in enumerate(notes, start=1):
            # Get token text for the note
            token_text = self._get_note_token_text(note, sentence, token_id_to_order)

            # Format note: "1. "quoted tokens" in italics - note text"
            if token_text:
                note_line = f'{note_idx}. "{token_text}" - {note.note_text_md}'
            else:
                note_line = f"{note_idx}. {note.note_text_md}"

            doc.add_paragraph(note_line)

    def _sort_notes_by_position(self, sentence: Sentence) -> list:
        """
        Sort notes by their position in the sentence (by start token order_index).

        Args:
            sentence: Sentence to get notes from

        Returns:
            Sorted list of notes

        """
        # Build token ID to order_index mapping
        token_id_to_order: dict[int, int] = {}
        for token in sentence.tokens:
            if token.id:
                token_id_to_order[token.id] = token.order_index

        def get_note_position(note) -> int:
            """Get position of note in sentence based on start token."""
            if note.start_token and note.start_token in token_id_to_order:
                return token_id_to_order[note.start_token]
            # Fallback to end_token if start_token not found
            if note.end_token and note.end_token in token_id_to_order:
                return token_id_to_order[note.end_token]
            # Fallback to very high number if neither found
            return 999999

        # Sort notes by position
        return sorted(sentence.notes, key=get_note_position)

    def _get_note_token_text(
        self,
        note,
        sentence: Sentence,
        token_id_to_order: dict[int, int],  # noqa: ARG002
    ) -> str:
        """
        Get token text for a note.

        Args:
            note: Note to get tokens for
            sentence: Sentence containing the tokens
            token_id_to_order: Map of token ID to order_index

        Returns:
            Token text string (space-separated tokens)

        """
        if not note.start_token or not note.end_token:
            return ""

        # Find tokens by ID
        start_token = None
        end_token = None
        for token in sentence.tokens:
            if token.id == note.start_token:
                start_token = token
            if token.id == note.end_token:
                end_token = token

        if not start_token or not end_token:
            return ""

        # Get all tokens in range
        tokens_in_range = []
        in_range = False
        for token in sorted(sentence.tokens, key=lambda t: t.order_index):
            if token.id == start_token.id:
                in_range = True
            if in_range:
                tokens_in_range.append(token.surface)
            if token.id == end_token.id:
                break

        return " ".join(tokens_in_range)

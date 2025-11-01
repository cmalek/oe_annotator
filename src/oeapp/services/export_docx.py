"""DOCX export service for Old English Annotator."""

from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor

from src.oeapp.services.db import Database
from src.oeapp.models.annotation import Annotation


class DOCXExporter:
    """Exports annotated Old English text to DOCX format."""

    def __init__(self, db: Database):
        """Initialize exporter.

        Args:
            db: Database connection
        """
        self.db = db

    def export(self, project_id: int, output_path: Path) -> bool:
        """Export project to DOCX file.

        Args:
            project_id: Project ID to export
            output_path: Path to output DOCX file

        Returns:
            True if successful, False otherwise
        """
        try:
            doc = Document()
            self._setup_document_styles(doc)

            cursor = self.db.conn.cursor()

            # Get project name
            cursor.execute("SELECT name FROM projects WHERE id = ?", (project_id,))
            project_row = cursor.fetchone()
            project_name = project_row["name"] if project_row else "Untitled Project"

            # Add title
            title = doc.add_heading(project_name, level=1)
            doc.add_paragraph()  # Blank line after title

            # Get sentences in order
            cursor.execute(
                "SELECT id, display_order, text_oe, text_modern FROM sentences "
                "WHERE project_id = ? ORDER BY display_order",
                (project_id,)
            )
            sentences = cursor.fetchall()

            for sentence_row in sentences:
                sentence_id = sentence_row["id"]
                display_order = sentence_row["display_order"]
                text_oe = sentence_row["text_oe"]
                text_modern = sentence_row["text_modern"]

                # Add sentence number
                sentence_num_para = doc.add_paragraph()
                sentence_num_run = sentence_num_para.add_run(f"[{display_order}] ")
                sentence_num_run.bold = True

                # Build Old English sentence with annotations
                self._add_oe_sentence_with_annotations(doc, sentence_id, text_oe)

                # Add translation
                if text_modern:
                    trans_para = doc.add_paragraph(text_modern)
                else:
                    # Empty translation paragraph
                    doc.add_paragraph()

                # Blank line
                doc.add_paragraph()

                # Add notes
                self._add_notes(doc, sentence_id)

                # Blank line between sentences
                doc.add_paragraph()

            doc.save(str(output_path))
            return True
        except Exception as e:
            print(f"Export error: {e}")
            return False

    def _setup_document_styles(self, doc: Document):
        """Set up document styles.

        Args:
            doc: Document to set up
        """
        # Styles are already available in python-docx
        # Title, Body, Default are standard styles
        pass

    def _add_oe_sentence_with_annotations(self, doc: Document, sentence_id: int, text_oe: str):
        """Add Old English sentence with superscript/subscript annotations.

        Args:
            doc: Document to add to
            sentence_id: Sentence ID
            text_oe: Old English text
        """
        cursor = self.db.conn.cursor()

        # Get tokens for this sentence
        cursor.execute(
            "SELECT id, order_index, surface FROM tokens WHERE sentence_id = ? ORDER BY order_index",
            (sentence_id,)
        )
        tokens = cursor.fetchall()

        # Get annotations for all tokens
        token_ids = [t["id"] for t in tokens]
        if not token_ids:
            # No tokens, just add plain text
            para = doc.add_paragraph(text_oe)
            for run in para.runs:
                run.italic = True
            return

        annotations = {}
        cursor.execute(
            "SELECT * FROM annotations WHERE token_id IN ({})".format(
                ",".join(["?"] * len(token_ids))
            ),
            token_ids
        )
        for row in cursor.fetchall():
            ann = Annotation(
                token_id=row["token_id"],
                pos=row["pos"],
                gender=row["gender"],
                number=row["number"],
                case=row["case"],
                declension=row["declension"],
                pronoun_type=row["pronoun_type"],
                verb_class=row["verb_class"],
                verb_tense=row["verb_tense"],
                verb_person=row["verb_person"],
                verb_mood=row["verb_mood"],
                verb_aspect=row["verb_aspect"],
                verb_form=row["verb_form"],
                prep_case=row["prep_case"],
                uncertain=bool(row["uncertain"]),
                alternatives_json=row["alternatives_json"],
                confidence=row["confidence"],
            )
            annotations[row["token_id"]] = ann

        # Build paragraph with annotations
        para = doc.add_paragraph()

        for i, token in enumerate(tokens):
            token_id = token["id"]
            surface = token["surface"]
            annotation = annotations.get(token_id)

            # Add word
            word_run = para.add_run(surface)
            word_run.italic = True

            if annotation:
                # Build superscript text (POS abbreviations, case, number, gender)
                superscript_parts = []

                # Add POS abbreviations first
                pos_label = self._format_pos(annotation)
                if pos_label:
                    superscript_parts.append(pos_label)

                # Add case/number/gender compact format
                if annotation.case and annotation.number:
                    case_map = {"n": "n", "a": "a", "g": "g", "d": "d", "i": "i"}
                    number_map = {"s": "1", "p": "pl"}
                    case_str = case_map.get(annotation.case, "")
                    number_str = number_map.get(annotation.number, "")
                    if case_str and number_str:
                        compact = f"{case_str}{number_str}"
                        superscript_parts.append(compact)
                elif annotation.case:
                    case_map = {"n": "n", "a": "a", "g": "g", "d": "d", "i": "i"}
                    superscript_parts.append(case_map.get(annotation.case, annotation.case))
                elif annotation.number:
                    number_map = {"s": "1", "p": "pl"}
                    superscript_parts.append(number_map.get(annotation.number, annotation.number))
                if annotation.gender:
                    superscript_parts.append(annotation.gender)

                # Build subscript text (detailed morphological info)
                subscript_parts = []
                if annotation.declension:
                    subscript_parts.append(annotation.declension)
                if annotation.verb_class:
                    subscript_parts.append(annotation.verb_class)
                # Add compact case/number combinations like "dat1", "acc1"
                if annotation.case and annotation.number:
                    case_map = {"n": "nom", "a": "acc", "g": "gen", "d": "dat", "i": "inst"}
                    number_map = {"s": "1", "p": "pl"}
                    case_str = case_map.get(annotation.case, annotation.case)[:3]
                    number_str = number_map.get(annotation.number, annotation.number)
                    if annotation.case in ["d", "a", "g"]:  # Common cases for compact format
                        compact = f"{case_str}{number_str}"
                        if compact not in subscript_parts:
                            subscript_parts.insert(0, compact)

                # Add uncertain marker
                uncertain_marker = "?" if annotation.uncertain else ""
                if annotation.alternatives_json:
                    alternatives = f" / {annotation.alternatives_json}"
                else:
                    alternatives = ""

                # Add superscript
                if superscript_parts:
                    sup_text = "".join(superscript_parts) + uncertain_marker + alternatives
                    sup_run = para.add_run(sup_text)
                    sup_run.font.size = Pt(8)
                    sup_run.font.superscript = True

                # Add subscript
                if subscript_parts:
                    sub_text = "".join(subscript_parts) + uncertain_marker
                    sub_run = para.add_run(sub_text)
                    sub_run.font.size = Pt(8)
                    sub_run.font.subscript = True

            # Add space after word (except last token)
            if i < len(tokens) - 1:
                para.add_run(" ")

    def _format_pos(self, annotation: Annotation) -> str:
        """Format POS abbreviation for display.

        Args:
            annotation: Annotation object

        Returns:
            Formatted POS string
        """
        if not annotation.pos:
            return ""

        pos_map = {
            "N": "n:",
            "V": "v:",
            "A": "adj:",
            "R": "pron:",
        }

        base = pos_map.get(annotation.pos, annotation.pos.lower())

        if annotation.pos == "R" and annotation.pronoun_type:
            type_map = {"p": "pers", "r": "rel", "d": "dem", "i": "int"}
            type_str = type_map.get(annotation.pronoun_type, "")
            if type_str:
                return f"{base}{type_str}"
        elif annotation.pos == "V" and annotation.verb_class:
            return f"{base}{annotation.verb_class}"

        return base

    def _add_notes(self, doc: Document, sentence_id: int):
        """Add notes for a sentence.

        Args:
            doc: Document to add to
            sentence_id: Sentence ID
        """
        cursor = self.db.conn.cursor()

        # Get notes for this sentence
        cursor.execute(
            "SELECT start_token, note_text_md FROM notes "
            "WHERE sentence_id = ? AND note_type IN ('token', 'span', 'sentence') "
            "ORDER BY start_token, id",
            (sentence_id,)
        )
        notes = cursor.fetchall()

        if not notes:
            return

        # Get token surfaces for note references
        cursor.execute(
            "SELECT id, surface FROM tokens WHERE sentence_id = ? ORDER BY order_index",
            (sentence_id,)
        )
        tokens_dict = {row["id"]: row["surface"] for row in cursor.fetchall()}

        # Add notes as numbered list
        for i, note_row in enumerate(notes, 1):
            token_id = note_row["start_token"]
            note_text = note_row["note_text_md"]

            # Get token surface if available
            token_surface = tokens_dict.get(token_id, "") if token_id else ""

            # Format note
            if token_surface:
                note_line = f"{token_surface}, {note_text}"
            else:
                note_line = note_text

            para = doc.add_paragraph(note_line)
            # Could add numbering here if needed
